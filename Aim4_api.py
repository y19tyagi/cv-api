from flask import Flask, request, jsonify
from docx import Document
import base64
import io
import zipfile
from PIL import Image

app = Flask(__name__)

@app.route('/extract', methods=['POST'])
def extract():
    try:
        # Get base64 from request
        base64_str = request.json.get('base64')
        if not base64_str:
            return jsonify({"error": "Missing base64 field"}), 400

        # Decode base64 to bytes
        docx_bytes = base64.b64decode(base64_str)
        file_stream = io.BytesIO(docx_bytes)

        # --- Extract text ---
        document = Document(file_stream)

        # 🔁 Recursive table extractor
        def extract_text_from_tables(tables):
            texts = []
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text.strip())
                        if cell.tables:
                            texts.extend(extract_text_from_tables(cell.tables))
            return texts

        # 📄 Combine paragraphs and all tables
        text_parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        text_parts.extend(extract_text_from_tables(document.tables))
        text = "\n".join(text_parts)

        # --- Extract images ---
        images = []
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as docx_zip:
            for file in docx_zip.namelist():
                if file.startswith("word/media/") and file.lower().endswith((".png", ".jpg", ".jpeg")):
                    img_data = docx_zip.read(file)
                    img_base64 = base64.b64encode(img_data).decode('utf-8')
                    images.append(img_base64)

        return jsonify({
            "text": text,
            "photos": images
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)
